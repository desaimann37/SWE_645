from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Inches(1.15)
section.right_margin  = Inches(1.15)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Default font ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

# ── Helper: set cell background color ─────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def set_cell_borders(table, color="CCCCCC"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), color)
                tcBorders.append(border)
            tcPr.append(tcBorders)

# ── Helper: heading 2 with bottom border ─────────────────────────────────
def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"),   "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"),  "F5F5F5")
    p._p.get_or_add_pPr().append(shading)

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(10)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(8.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

def add_image(doc, path, width_in, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    add_caption(doc, caption)

def mk_table(doc, headers, rows, col_widths_in, alt_color="F4F6F8"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Inches(col_widths_in[i])
        set_cell_bg(cell, "2C3E50")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else alt_color
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Inches(col_widths_in[ci])
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9.5)

    set_cell_borders(t)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(72)
r = cover.add_run("Design Document")
r.bold = True
r.font.size = Pt(22)
r.font.name = "Times New Roman"
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Agentic AI Extension of the Student Survey System")
r2.font.size = Pt(13)
r2.font.name = "Times New Roman"
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Using LangGraph and FastMCP")
r3.font.size = Pt(13)
r3.font.name = "Times New Roman"
r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

for line in [
    "Course: SWE 645 — Extra Credit Assignment",
    "Team: Mann Desai  |  Tisha  |  Aakash  |  Yash",
    "Date: April 2026",
]:
    pm = doc.add_paragraph(line)
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pm.paragraph_format.space_after = Pt(4)
    for run in pm.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SECTION 1 — Introduction
# ══════════════════════════════════════════════════════════════════════════
add_h2(doc, "1. Introduction")

add_body(doc,
    "For this extra credit assignment, we took the Student Survey system we built in "
    "Homework 3 and added an AI-powered layer on top of it. The original system let users "
    "submit and manage surveys through a standard React form and a FastAPI REST backend. "
    "What we added here is a completely separate conversational interface so that a user "
    "can type something like \"delete John Doe's survey\" and the system will figure out "
    "what they mean, find the right record, ask for confirmation, and carry out the action "
    "— all without the user clicking through any forms."
)
add_body(doc,
    "The core idea was to keep everything we already had intact and build the AI "
    "orchestration layer on top of it. We didn't rewrite the backend or touch the existing "
    "REST routes; we just exposed the same database operations as MCP tools so the agent "
    "could call them."
)
add_body(doc, "The overall architecture is shown in Figure 1 below:")

add_image(doc, "arch-diagram.png", 5.8,
    "Figure 1 — System Architecture: three Kubernetes pods with LangGraph agent, "
    "FastMCP backend, and MySQL RDS.")

add_body(doc,
    "Each of those three pods is independently deployed on Kubernetes and managed via Helm "
    "charts. The agent communicates with the MCP backend over HTTP inside the cluster, and "
    "the React frontend talks to the agent through a simple REST endpoint exposed at "
    "POST /agent/query."
)


# ══════════════════════════════════════════════════════════════════════════
# SECTION 2 — LangGraph Graph Design
# ══════════════════════════════════════════════════════════════════════════
add_h2(doc, "2. LangGraph Graph Design")

add_body(doc,
    "The agent is built as a LangGraph StateGraph. We chose LangGraph because it lets you "
    "define exactly what happens at each step and how decisions are made, rather than "
    "relying on the LLM to figure all of that out on its own. For a system that can delete "
    "database records, that level of control matters."
)

add_h3(doc, "State Object")
add_body(doc,
    "All the information the agent needs across a conversation is stored in a typed "
    "dictionary called AgentState. The most important fields are:"
)

mk_table(doc,
    ["Field", "What it holds"],
    [
        ["messages",         "Full chat history so the LLM has context"],
        ["intent",           "What the user wants: create / read / update / delete / confirm / cancel"],
        ["draft_survey",     "Partially filled survey during a multi-turn create flow"],
        ["missing_fields",   "Which fields the agent still needs to ask about"],
        ["target_survey_id", "ID of the survey being updated or deleted"],
        ["update_changes",   "Specific fields to change in an update"],
        ["awaiting",         "Whether we're waiting for user confirmation before executing"],
        ["response_text",    "The final message to send back to the user"],
    ],
    col_widths_in=[1.9, 3.8],
)

add_body(doc,
    "The awaiting field is what makes the human-in-the-loop safety work. When the agent "
    "identifies a survey to delete and displays its details, it sets "
    "awaiting = \"delete_confirm\" and returns to the user. On the next message, the router "
    "checks that field before doing anything — if the user said \"yes,\" it routes to "
    "delete_execute; otherwise it cancels."
)

add_h3(doc, "Graph Topology")
add_body(doc, "The full LangGraph workflow is shown in Figure 2:")

add_image(doc, "langgraph-diagram.png", 5.8,
    "Figure 2 — LangGraph workflow: intent classification routes into four CRUD branches, "
    "each with a confirmation gate before destructive operations.")

add_h3(doc, "How Routing Works")
add_body(doc,
    "After the intent node runs, a conditional edge function called _route_from_intent "
    "decides where to go next. It first checks whether awaiting is set — if the user just "
    "replied \"yes\" to a pending delete confirmation, it routes straight to delete_execute "
    "regardless of how the LLM interpreted the word \"yes.\" That check takes priority over "
    "everything else. Only if there's no pending confirmation does it route based on the "
    "classified intent."
)

add_h3(doc, "Session Persistence")
add_body(doc,
    "LangGraph graphs don't carry state between HTTP requests by themselves. We handle "
    "this with an in-memory Python dictionary keyed by session_id. Every time the user "
    "sends a message, the server rehydrates the stored state, runs the graph, and saves "
    "the updated state back. It's simple and works well for a per-session conversation "
    "without needing a full external state store."
)


# ══════════════════════════════════════════════════════════════════════════
# SECTION 3 — MCP Tools
# ══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h2(doc, "3. MCP Tools Implemented")

add_body(doc,
    "The tool layer lives in backend/mcp_tools.py and is served by FastMCP, which is "
    "mounted as a sub-app inside our existing FastAPI backend at the path /mcp. The agent "
    "connects to it using langchain-mcp-adapters. We implemented all six tools the "
    "assignment required:"
)

mk_table(doc,
    ["Tool", "Parameters", "What it does"],
    [
        ["create_survey",    "All 12 required fields",         "Inserts a new survey and returns it with its assigned ID"],
        ["list_surveys",     "None",                            "Returns every survey record in the database"],
        ["get_survey_by_id", "survey_id",                      "Fetches a single survey by its numeric ID"],
        ["search_surveys",   "9 optional filters",             "Case-insensitive partial-match search by name, city, state, liked_most, interested_via, recommendation, and date range"],
        ["update_survey",    "survey_id + any subset of fields","Patches only the provided fields, leaves others unchanged"],
        ["delete_survey",    "survey_id",                      "Deletes the survey and returns a snapshot of what was deleted"],
    ],
    col_widths_in=[1.5, 1.5, 2.7],
)

add_body(doc,
    "All tools reuse the same SQLModel ORM session that the REST API uses, so there's "
    "one source of truth for the database schema. Each tool has a docstring describing its "
    "behavior — this matters because the LLM uses those descriptions to decide which "
    "tool to call."
)


# ══════════════════════════════════════════════════════════════════════════
# SECTION 4 — React Chat Interface
# ══════════════════════════════════════════════════════════════════════════
add_h2(doc, "4. React Chat Interface")

add_body(doc,
    "We added a new page called \"AI Survey Assistant\" to the existing React app, "
    "accessible from a new nav link. The layout follows the suggested design from the "
    "assignment spec, with three panels:"
)

bullets = [
    ("Left panel",  "explains what the agent can do, lists all the fields needed to create a survey, "
                    "and includes eight clickable example prompts users can try without having to type anything."),
    ("Main chat area", "a scrollable message thread with user messages on the right (dark blue bubbles) "
                       "and agent responses on the left (white bubbles). It shows a \"Thinking...\" indicator "
                       "while waiting for the agent and supports Enter-to-send."),
    ("Right panel", "a small log of recent agent actions (created, updated, deleted, queried) so the user "
                    "can track what has been done without scrolling back through the conversation."),
]
for label, text in bullets:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    bold_run = p.add_run(label + " — ")
    bold_run.bold = True
    bold_run.font.name = "Times New Roman"
    bold_run.font.size = Pt(11)
    norm_run = p.add_run(text)
    norm_run.font.name = "Times New Roman"
    norm_run.font.size = Pt(11)

doc.add_paragraph()
add_body(doc,
    "The component generates a UUID session ID when it first loads and sends it with "
    "every message. This is how the server knows which in-memory state to load for that "
    "conversation. Multi-turn state — like remembering we're halfway through collecting "
    "fields for a new survey — is handled entirely on the server side. The frontend just "
    "sends messages and renders responses."
)


# ══════════════════════════════════════════════════════════════════════════
# SECTION 5 — Kubernetes / Helm
# ══════════════════════════════════════════════════════════════════════════
add_h2(doc, "5. Kubernetes and Helm Deployment")

add_body(doc,
    "We created a Helm chart at helm/student-survey/ that deploys the full system as "
    "three separate pods, matching exactly what the assignment requires:"
)

mk_table(doc,
    ["Pod", "Image", "Service Type"],
    [
        ["Frontend (React + Nginx)", "student-survey-frontend", "LoadBalancer — port 80"],
        ["Agent (LangGraph)",        "student-survey-agent",    "ClusterIP — port 9000"],
        ["Backend (FastAPI + MCP)",  "student-survey-backend",  "ClusterIP — port 8000"],
    ],
    col_widths_in=[1.9, 2.1, 1.7],
)

add_body(doc,
    "Sensitive values — the Anthropic API key and the database connection string — are "
    "stored as Kubernetes Secrets and injected as environment variables at pod startup. "
    "No credentials are hardcoded anywhere in the Helm chart templates. The values.yaml "
    "has placeholder fields for base64-encoded secrets that get passed in at deploy time "
    "via --set."
)
add_body(doc,
    "The agent deployment reads MCP_SERVER_URL pointing to the backend service's "
    "cluster-internal DNS name, so all communication between the agent and the backend "
    "stays inside the cluster. Only the frontend is exposed externally through the "
    "LoadBalancer."
)


# ══════════════════════════════════════════════════════════════════════════
# SECTION 6 — Challenges
# ══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h2(doc, "6. Challenges We Ran Into")

challenges = [
    (
        "MCP Tool Result Format",
        "This one took a while to figure out. Depending on the version of "
        "langchain-mcp-adapters, the same tool call could return a plain Python dict, "
        "a JSON string, or a list of MCP content blocks that looked like "
        "[{\"type\": \"text\", \"text\": \"...\"}]. The agent nodes were breaking because "
        "they expected a dict but kept getting a list. We ended up writing a small "
        "_unwrap_tool_result() utility that normalizes all three formats before any node "
        "tries to read from the result."
    ),
    (
        "Keeping Confirmation State Across HTTP Requests",
        "The trickiest part of the whole project was making multi-turn confirmation work "
        "reliably. When the agent says \"I found John Doe's survey. Do you want to delete "
        "it?\" and the user replies \"yes,\" the server receives a brand new HTTP request "
        "with no built-in memory. If we just ran intent classification on \"yes,\" it would "
        "come back as confirm — but the router also needs to know it's confirming a delete, "
        "not a create or update. The fix was to save the awaiting field in the session store "
        "and check it before routing, so the router always knows what confirmation is pending."
    ),
    (
        "FastMCP Lifespan Inside FastAPI",
        "FastMCP requires its own lifespan context to initialize the MCP session manager. "
        "When we mounted it as a sub-app inside FastAPI, the session manager wasn't starting "
        "and every tool call was failing with a connection error. The fix was to nest the "
        "MCP lifespan inside our FastAPI lifespan using asynccontextmanager, so both start "
        "up together. Once we figured that out it worked perfectly, but it wasn't obvious "
        "from the FastMCP docs."
    ),
    (
        "Intent Classification for Short Replies",
        "Early on, when a user typed \"yes\" after a delete confirmation prompt, the LLM "
        "would sometimes classify it as unknown or even create depending on conversation "
        "context. We fixed this by adding confirm and cancel as explicit intent classes in "
        "the classification prompt with worked examples, and by always routing based on the "
        "awaiting field first before trusting the classifier."
    ),
    (
        "Vite Environment Variables at Build Time",
        "We wanted to pass the agent URL into the React app via Docker Compose, but Vite "
        "bakes VITE_* variables at image build time, not at container runtime. So environment "
        "variables set in docker-compose.yaml under the frontend service don't actually get "
        "picked up. We resolved this by setting safe fallback values in the component code "
        "(http://localhost:8000 and http://localhost:9000), which work correctly when running "
        "locally since those ports are exposed on the host."
    ),
]

for title, body in challenges:
    add_h3(doc, title)
    add_body(doc, body)


# ══════════════════════════════════════════════════════════════════════════
# SECTION 7 — Tech Stack
# ══════════════════════════════════════════════════════════════════════════
add_h2(doc, "7. Tech Stack Summary")

mk_table(doc,
    ["Layer", "Technology"],
    [
        ["Frontend",      "React 19, React Router, Axios, Vite, Nginx"],
        ["AI Agent",      "LangGraph, LangChain, Claude Haiku 4.5 (Anthropic), FastAPI"],
        ["Tool Server",   "FastMCP 2.10+, langchain-mcp-adapters"],
        ["Backend ORM",   "SQLModel (SQLAlchemy + Pydantic)"],
        ["Database",      "MySQL 8 on AWS RDS"],
        ["Containers",    "Docker, Docker Compose"],
        ["Orchestration", "Kubernetes, Helm 3"],
    ],
    col_widths_in=[1.8, 3.9],
)


# ══════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════
doc.save("design-document.docx")
print("design-document.docx saved")
